from rest_framework import viewsets, status, generics
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django_filters.rest_framework import DjangoFilterBackend
from django.http import HttpResponse
from django.utils import timezone
from io import BytesIO
from django.db.models import Count

from .models import ReportTemplate, GeneratedReport
from .serializers import ReportTemplateSerializer, GeneratedReportSerializer
from apps.accounts.models import AuditTrail


class ReportTemplateViewSet(viewsets.ModelViewSet):
    queryset = ReportTemplate.objects.all()
    serializer_class = ReportTemplateSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['template_type', 'is_default']

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)


class GeneratedReportViewSet(viewsets.ModelViewSet):
    queryset = GeneratedReport.objects.select_related(
        'template', 'engagement', 'generated_by'
    ).all()
    serializer_class = GeneratedReportSerializer
    permission_classes = [IsAuthenticated]
    filter_backends = [DjangoFilterBackend]
    filterset_fields = ['format', 'status', 'engagement']

    def perform_create(self, serializer):
        report = serializer.save(generated_by=self.request.user, status='generating')
        try:
            self.generate_report_file(report)
        except Exception as e:
            report.status = 'failed'
            report.error_message = str(e)
            report.save()

    def generate_report_file(self, report):
        from django.core.files.base import ContentFile
        from apps.findings.models import AuditFinding
        from apps.corrective_actions.models import CorrectiveAction
        from apps.risk_assessment.models import RiskAssessment
        
        findings = []
        corrective_actions = []
        risk_data = []
        engagement = report.engagement
        engagement_info = {}
        
        if engagement:
            findings = list(AuditFinding.objects.filter(engagement=engagement))
            corrective_actions = list(CorrectiveAction.objects.filter(
                finding__engagement=engagement
            ).select_related('finding'))
            risk_data = list(RiskAssessment.objects.filter(
                department=engagement.department
            ) if engagement.department else [])
            
            engagement_info = {
                'title': engagement.title,
                'number': engagement.engagement_number,
                'type': engagement.get_engagement_type_display(),
                'department': engagement.department.name if engagement.department else 'N/A',
                'status': engagement.get_status_display(),
                'objectives': engagement.objectives,
                'scope': engagement.scope,
                'lead_auditor': engagement.lead_auditor.full_name if engagement.lead_auditor else 'N/A',
                'supervisor': engagement.supervisor.full_name if engagement.supervisor else 'N/A',
                'planned_start': engagement.planned_start,
                'planned_end': engagement.planned_end,
                'actual_start': engagement.actual_start,
                'actual_end': engagement.actual_end,
                'risk_level': engagement.risk_level,
            }
            
        buf = BytesIO()
        filename = f"{report.title.replace(' ', '_')}.{report.format}"

        if report.format == 'pdf':
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            from reportlab.lib.units import mm
            from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
            
            doc = SimpleDocTemplate(buf, pagesize=A4,
                                    leftMargin=20*mm, rightMargin=20*mm,
                                    topMargin=20*mm, bottomMargin=20*mm)
            styles = getSampleStyleSheet()
            
            # Create custom styles
            styles.add(ParagraphStyle('CoverTitle', parent=styles['Title'],
                                      fontSize=22, spaceAfter=6, alignment=TA_CENTER,
                                      textColor=colors.HexColor('#1E3A5F')))
            styles.add(ParagraphStyle('CoverSubtitle', parent=styles['Normal'],
                                      fontSize=14, alignment=TA_CENTER,
                                      textColor=colors.HexColor('#555555')))
            styles.add(ParagraphStyle('SectionTitle', parent=styles['Heading1'],
                                      fontSize=16, spaceBefore=20, spaceAfter=10,
                                      textColor=colors.HexColor('#1E3A5F'),
                                      borderWidth=2, borderColor=colors.HexColor('#1E3A5F'),
                                      borderPadding=4))
            styles.add(ParagraphStyle('SectionBody', parent=styles['Normal'],
                                      fontSize=10, leading=14, alignment=TA_JUSTIFY,
                                      spaceAfter=8))
            styles.add(ParagraphStyle('TableHeader', parent=styles['Normal'],
                                      fontSize=9, textColor=colors.white))
            
            elements = []
            
            # ========== COVER PAGE ==========
            elements.append(Spacer(1, 60))
            elements.append(Paragraph("ETHIOPIAN ELECTRIC UTILITY", styles['CoverTitle']))
            elements.append(Paragraph("Internal Audit Department", styles['CoverSubtitle']))
            elements.append(Spacer(1, 20))
            elements.append(Paragraph(report.title, styles['CoverTitle']))
            elements.append(Spacer(1, 30))
            
            # Engagement info box
            info_data = [
                ['Engagement:', engagement_info.get('title', 'N/A')],
                ['Eng. Number:', engagement_info.get('number', 'N/A')],
                ['Type:', engagement_info.get('type', 'N/A')],
                ['Department:', engagement_info.get('department', 'N/A')],
                ['Lead Auditor:', engagement_info.get('lead_auditor', 'N/A')],
                ['Supervisor:', engagement_info.get('supervisor', 'N/A')],
                ['Risk Level:', engagement_info.get('risk_level', 'N/A')],
            ]
            info_table = Table(info_data, colWidths=[100, 320])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 40))
            elements.append(Paragraph(f"Date: {timezone.now().strftime('%d %B %Y')}", styles['CoverSubtitle']))
            elements.append(Paragraph(f"Generated By: {report.generated_by.full_name if report.generated_by else 'System'}", styles['CoverSubtitle']))
            elements.append(PageBreak())
            
            # ========== TABLE OF CONTENTS ==========
            elements.append(Paragraph("TABLE OF CONTENTS", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            toc_items = [
                "1.  Background",
                "2.  Executive Summary",
                "3.  Risk Analysis",
                "4.  Findings Summary Table",
                "5.  Corrective Action Plan (CAPA)",
            ]
            for item in toc_items:
                elements.append(Paragraph(item, styles['SectionBody']))
            elements.append(PageBreak())
            
            # ========== 1. BACKGROUND ==========
            elements.append(Paragraph("1. BACKGROUND", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            elements.append(Paragraph(
                f"This report presents the findings and recommendations from the audit engagement "
                f"<b>{engagement_info.get('title', 'N/A')}</b> conducted by the EEU Internal Audit Department. "
                f"The audit was performed in accordance with the International Standards for the Professional "
                f"Practice of Internal Auditing (IPPF) and the EEU Internal Audit Charter.",
                styles['SectionBody']
            ))
            elements.append(Spacer(1, 10))
            
            bg_data = [
                ['Engagement Title', engagement_info.get('title', 'N/A')],
                ['Engagement Type', engagement_info.get('type', 'N/A')],
                ['Audited Department', engagement_info.get('department', 'N/A')],
                ['Status', engagement_info.get('status', 'N/A')],
                ['Lead Auditor', engagement_info.get('lead_auditor', 'N/A')],
                ['Supervisor', engagement_info.get('supervisor', 'N/A')],
                ['Planned Start', str(engagement_info.get('planned_start', 'N/A'))],
                ['Planned End', str(engagement_info.get('planned_end', 'N/A'))],
            ]
            bg_table = Table(bg_data, colWidths=[140, 280])
            bg_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('TOPPADDING', (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0F4F8')),
            ]))
            elements.append(bg_table)
            elements.append(Spacer(1, 15))
            
            objectives_text = engagement_info.get('objectives', 'No specific objectives defined for this engagement.')
            scope_text = engagement_info.get('scope', 'No scope defined for this engagement.')
            elements.append(Paragraph("<b>Audit Objectives:</b>", styles['SectionBody']))
            elements.append(Paragraph(objectives_text, styles['SectionBody']))
            elements.append(Spacer(1, 8))
            elements.append(Paragraph("<b>Audit Scope:</b>", styles['SectionBody']))
            elements.append(Paragraph(scope_text, styles['SectionBody']))
            elements.append(PageBreak())
            
            # ========== 2. EXECUTIVE SUMMARY ==========
            elements.append(Paragraph("2. EXECUTIVE SUMMARY", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            
            total_findings = len(findings)
            severity_counts = {}
            status_counts = {}
            for f in findings:
                sev = f.severity.upper() if f.severity else 'UNKNOWN'
                severity_counts[sev] = severity_counts.get(sev, 0) + 1
                sts = f.status.upper() if f.status else 'UNKNOWN'
                status_counts[sts] = status_counts.get(sts, 0) + 1
            
            critical_count = severity_counts.get('CRITICAL', 0)
            high_count = severity_counts.get('HIGH', 0)
            medium_count = severity_counts.get('MEDIUM', 0)
            low_count = severity_counts.get('LOW', 0)
            
            elements.append(Paragraph(
                f"A total of <b>{total_findings}</b> finding(s) were identified during this audit engagement. "
                f"Of these, <b>{critical_count}</b> are classified as Critical, <b>{high_count}</b> as High, "
                f"<b>{medium_count}</b> as Medium, and <b>{low_count}</b> as Low severity. "
                f"The audit was conducted to evaluate the adequacy and effectiveness of internal controls "
                f"within the audited area.",
                styles['SectionBody']
            ))
            elements.append(Spacer(1, 10))
            
            if total_findings > 0:
                exec_data = [['Severity Level', 'Count']]
                for sev, count in sorted(severity_counts.items(), key=lambda x: ['CRITICAL','HIGH','MEDIUM','LOW'].index(x[0]) if x[0] in ['CRITICAL','HIGH','MEDIUM','LOW'] else 99):
                    exec_data.append([sev, str(count)])
                
                exec_table = Table(exec_data, colWidths=[200, 200])
                exec_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('ALIGN', (1, 0), (1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#1E3A5F')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ]))
                elements.append(exec_table)
            elements.append(PageBreak())
            
            # ========== 3. RISK ANALYSIS ==========
            elements.append(Paragraph("3. RISK ANALYSIS", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            
            if risk_data:
                elements.append(Paragraph(
                    f"The following risk assessment data relates to the department <b>{engagement_info.get('department', 'N/A')}</b>. "
                    f"The risk analysis considers inherent risk, control effectiveness, and residual risk levels.",
                    styles['SectionBody']
                ))
                elements.append(Spacer(1, 10))
                
                risk_table_data = [['Period', 'Year', 'Likelihood', 'Impact', 'Risk Score', 'Rating', 'Control Eff.', 'Residual Risk']]
                for r in risk_data[:15]:
                    risk_table_data.append([
                        r.assessment_period,
                        str(r.year),
                        str(r.likelihood),
                        str(r.impact),
                        str(r.risk_score),
                        r.risk_rating.upper() if r.risk_rating else 'N/A',
                        str(r.control_effectiveness),
                        f"{r.residual_risk:.2f}" if r.residual_risk else 'N/A',
                    ])
                
                risk_table = Table(risk_table_data, colWidths=[45, 35, 50, 40, 55, 50, 55, 60])
                risk_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 7),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F8F9FA')),
                ]))
                elements.append(risk_table)
            else:
                elements.append(Paragraph(
                    "No risk assessment data is available for this engagement's department. "
                    "A risk analysis should be conducted as part of the audit planning process.",
                    styles['SectionBody']
                ))
            elements.append(PageBreak())
            
            # ========== 4. FINDINGS TABLE ==========
            elements.append(Paragraph("4. FINDINGS SUMMARY TABLE", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            
            if not findings:
                elements.append(Paragraph("No findings were registered for this engagement.", styles['SectionBody']))
            else:
                elements.append(Paragraph(
                    f"The following table provides a comprehensive summary of all <b>{total_findings}</b> finding(s) "
                    f"identified during this audit engagement.",
                    styles['SectionBody']
                ))
                elements.append(Spacer(1, 10))
                
                findings_table_data = [['#', 'Ref No.', 'Title', 'Severity', 'Category', 'Status', 'Recommendation']]
                for idx, f in enumerate(findings, 1):
                    sev_color = colors.HexColor('#F44336') if f.severity == 'critical' else \
                                colors.HexColor('#FF9800') if f.severity == 'high' else \
                                colors.HexColor('#FFC107') if f.severity == 'medium' else \
                                colors.HexColor('#4CAF50')
                    findings_table_data.append([
                        str(idx),
                        f.finding_number,
                        f.title[:50] + ('...' if len(f.title) > 50 else ''),
                        f.severity.upper() if f.severity else 'N/A',
                        f.category.replace('_', ' ').title() if f.category else 'N/A',
                        f.status.upper() if f.status else 'N/A',
                        f.recommendation[:60] + ('...' if len(f.recommendation) > 60 else '') if f.recommendation else 'N/A',
                    ])
                
                f_table = Table(findings_table_data, colWidths=[20, 60, 90, 60, 65, 55, 110])
                f_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 7),
                    ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                    ('ALIGN', (3, 0), (5, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F8F9FA')]),
                ]))
                elements.append(f_table)
                
                # Detailed findings section
                elements.append(Spacer(1, 20))
                elements.append(Paragraph("<b>Detailed Finding Descriptions:</b>", styles['SectionBody']))
                elements.append(Spacer(1, 10))
                for f in findings:
                    elements.append(Paragraph(
                        f"<b>{f.finding_number}: {f.title}</b> "
                        f"[Severity: {f.severity.upper() if f.severity else 'N/A'}] "
                        f"[Status: {f.status.upper() if f.status else 'N/A'}]",
                        styles['SectionBody']
                    ))
                    elements.append(Paragraph(f"<b>Condition:</b> {f.condition or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Criteria:</b> {f.criteria or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Cause:</b> {f.cause or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Effect/Impact:</b> {f.effect or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Recommendation:</b> {f.recommendation or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Management Response:</b> {f.management_response or 'N/A'}", styles['SectionBody']))
                    elements.append(Spacer(1, 8))
            elements.append(PageBreak())
            
            # ========== 5. CAPA - CORRECTIVE ACTION PLAN ==========
            elements.append(Paragraph("5. CORRECTIVE ACTION PLAN (CAPA)", styles['SectionTitle']))
            elements.append(Spacer(1, 10))
            
            if not corrective_actions:
                elements.append(Paragraph(
                    "No corrective actions have been assigned for findings in this engagement. "
                    "Corrective actions should be defined for each finding to address identified issues.",
                    styles['SectionBody']
                ))
            else:
                elements.append(Paragraph(
                    f"The following <b>{len(corrective_actions)}</b> corrective action(s) have been defined "
                    f"to address the findings identified during this audit engagement.",
                    styles['SectionBody']
                ))
                elements.append(Spacer(1, 10))
                
                capa_table_data = [['#', 'Action No.', 'Title', 'Finding Ref', 'Priority', 'Status', 'Owner', 'Due Date']]
                for idx, ca in enumerate(corrective_actions, 1):
                    capa_table_data.append([
                        str(idx),
                        ca.action_number,
                        ca.title[:45] + ('...' if len(ca.title) > 45 else ''),
                        ca.finding.finding_number if ca.finding else 'N/A',
                        ca.priority.upper() if ca.priority else 'N/A',
                        ca.status.upper() if ca.status else 'N/A',
                        ca.owner.full_name if ca.owner else 'Unassigned',
                        str(ca.due_date),
                    ])
                
                capa_table = Table(capa_table_data, colWidths=[18, 55, 85, 55, 48, 55, 60, 55])
                capa_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 7),
                    ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                    ('ALIGN', (3, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('TOPPADDING', (0, 0), (-1, -1), 4),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F8F9FA')]),
                ]))
                elements.append(capa_table)
                
                # Detailed CAPA descriptions
                elements.append(Spacer(1, 15))
                elements.append(Paragraph("<b>Detailed Corrective Action Descriptions:</b>", styles['SectionBody']))
                elements.append(Spacer(1, 10))
                for ca in corrective_actions:
                    elements.append(Paragraph(
                        f"<b>{ca.action_number}: {ca.title}</b> "
                        f"[Priority: {ca.priority.upper()} - Status: {ca.status.upper()}]",
                        styles['SectionBody']
                    ))
                    elements.append(Paragraph(f"<b>Finding:</b> {ca.finding.finding_number + ' - ' + ca.finding.title if ca.finding else 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Description:</b> {ca.description or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Recommendation:</b> {ca.recommendation or 'N/A'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Owner:</b> {ca.owner.full_name if ca.owner else 'Unassigned'}", styles['SectionBody']))
                    elements.append(Paragraph(f"<b>Due Date:</b> {ca.due_date}", styles['SectionBody']))
                    if ca.management_response:
                        elements.append(Paragraph(f"<b>Management Response:</b> {ca.management_response}", styles['SectionBody']))
                    elements.append(Spacer(1, 8))
            
            # Footer disclaimer
            elements.append(Spacer(1, 30))
            elements.append(Paragraph(
                "<i>This report was generated automatically by the EEU Internal Audit Management System. "
                "The information contained herein is confidential and intended solely for internal use "
                "by authorized personnel of Ethiopian Electric Utility.</i>",
                ParagraphStyle('Disclaimer', parent=styles['Normal'], fontSize=8,
                              textColor=colors.HexColor('#888888'), alignment=TA_CENTER)
            ))
            
            doc.build(elements)
            buf.seek(0)
            report.file.save(filename, ContentFile(buf.read()), save=False)
            report.status = 'ready'
            report.save()
            
            # Log the export
            AuditTrail.objects.create(
                user=report.generated_by,
                action='EXPORT',
                model_name='GeneratedReport',
                object_id=str(report.id),
                object_repr=f"Generated PDF report: {report.title}",
                ip_address=None,
            )

        elif report.format == 'excel':
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
            
            wb = openpyxl.Workbook()
            
            # =========== SHEET 1: COVER & BACKGROUND ===========
            ws_cover = wb.active
            ws_cover.title = "Cover & Background"
            ws_cover.sheet_properties.tabColor = "1E3A5F"
            
            header_fill = PatternFill(fill_type='solid', fgColor='1E3A5F')
            header_font = Font(bold=True, color='FFFFFF', size=11)
            title_font = Font(bold=True, size=16, color='1E3A5F')
            subtitle_font = Font(bold=True, size=12, color='1E3A5F')
            normal_font = Font(size=10)
            bold_font = Font(bold=True, size=10)
            
            thin_border = Border(
                left=Side(style='thin', color='CCCCCC'),
                right=Side(style='thin', color='CCCCCC'),
                top=Side(style='thin', color='CCCCCC'),
                bottom=Side(style='thin', color='CCCCCC'),
            )
            
            ws_cover.merge_cells('A1:H1')
            ws_cover['A1'] = 'ETHIOPIAN ELECTRIC UTILITY - Internal Audit Department'
            ws_cover['A1'].font = title_font
            ws_cover['A1'].alignment = Alignment(horizontal='center')
            
            ws_cover.merge_cells('A2:H2')
            ws_cover['A2'] = report.title
            ws_cover['A2'].font = Font(bold=True, size=14, color='1E3A5F')
            ws_cover['A2'].alignment = Alignment(horizontal='center')
            
            ws_cover['A4'] = f"Generated: {timezone.now().strftime('%d/%m/%Y %H:%M')}"
            ws_cover['A4'].font = normal_font
            ws_cover['A5'] = f"Generated By: {report.generated_by.full_name if report.generated_by else 'System'}"
            
            # Background details
            ws_cover['A7'] = '1. BACKGROUND INFORMATION'
            ws_cover['A7'].font = subtitle_font
            
            bg_fields = [
                ('Engagement Title:', engagement_info.get('title', 'N/A')),
                ('Engagement Number:', engagement_info.get('number', 'N/A')),
                ('Engagement Type:', engagement_info.get('type', 'N/A')),
                ('Department:', engagement_info.get('department', 'N/A')),
                ('Status:', engagement_info.get('status', 'N/A')),
                ('Lead Auditor:', engagement_info.get('lead_auditor', 'N/A')),
                ('Supervisor:', engagement_info.get('supervisor', 'N/A')),
                ('Risk Level:', engagement_info.get('risk_level', 'N/A')),
                ('Planned Start:', str(engagement_info.get('planned_start', 'N/A'))),
                ('Planned End:', str(engagement_info.get('planned_end', 'N/A'))),
            ]
            for i, (label, value) in enumerate(bg_fields):
                row = 8 + i
                ws_cover[f'A{row}'] = label
                ws_cover[f'A{row}'].font = bold_font
                ws_cover[f'B{row}'] = value
                ws_cover[f'B{row}'].font = normal_font
            
            row = 19
            ws_cover[f'A{row}'] = 'Audit Objectives:'
            ws_cover[f'A{row}'].font = bold_font
            ws_cover.merge_cells(f'A{row+1}:H{row+3}')
            ws_cover[f'A{row+1}'] = engagement_info.get('objectives', 'No objectives defined.')
            ws_cover[f'A{row+1}'].font = normal_font
            ws_cover[f'A{row+1}'].alignment = Alignment(wrap_text=True, vertical='top')
            
            ws_cover[f'A{row+5}'] = 'Audit Scope:'
            ws_cover[f'A{row+5}'].font = bold_font
            ws_cover.merge_cells(f'A{row+6}:H{row+8}')
            ws_cover[f'A{row+6}'] = engagement_info.get('scope', 'No scope defined.')
            ws_cover[f'A{row+6}'].font = normal_font
            ws_cover[f'A{row+6}'].alignment = Alignment(wrap_text=True, vertical='top')
            
            ws_cover.column_dimensions['A'].width = 25
            ws_cover.column_dimensions['B'].width = 50
            
            # =========== SHEET 2: EXECUTIVE SUMMARY ===========
            ws_exec = wb.create_sheet("Executive Summary")
            ws_exec.sheet_properties.tabColor = "4CAF50"
            
            ws_exec.merge_cells('A1:F1')
            ws_exec['A1'] = '2. EXECUTIVE SUMMARY'
            ws_exec['A1'].font = subtitle_font
            
            ws_exec['A3'] = f'Total Findings: {total_findings}'
            ws_exec['A3'].font = bold_font
            
            exec_headers = ['Severity Level', 'Count']
            for col_idx, text in enumerate(exec_headers, 1):
                cell = ws_exec.cell(row=5, column=col_idx, value=text)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center')
                cell.border = thin_border
            
            sev_order = ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL']
            row = 6
            for sev in sev_order:
                count = severity_counts.get(sev, 0)
                if count > 0:
                    ws_exec.cell(row=row, column=1, value=sev)
                    ws_exec.cell(row=row, column=2, value=count)
                    ws_exec.cell(row=row, column=1).font = bold_font
                    ws_exec.cell(row=row, column=2).alignment = Alignment(horizontal='center')
                    ws_exec.cell(row=row, column=2).font = normal_font
                    row += 1
            
            ws_exec.column_dimensions['A'].width = 25
            ws_exec.column_dimensions['B'].width = 15
            
            ws_exec[f'A{row+2}'] = 'Summary Description:'
            ws_exec[f'A{row+2}'].font = bold_font
            ws_exec.merge_cells(f'A{row+3}:F{row+5}')
            ws_exec[f'A{row+3}'] = (
                f"A total of {total_findings} finding(s) were identified during this audit engagement. "
                f"Of these, {critical_count} are Critical, {high_count} are High, "
                f"{medium_count} are Medium, and {low_count} are Low severity."
            )
            ws_exec[f'A{row+3}'].font = normal_font
            ws_exec[f'A{row+3}'].alignment = Alignment(wrap_text=True, vertical='top')
            
            # =========== SHEET 3: RISK ANALYSIS ===========
            ws_risk = wb.create_sheet("Risk Analysis")
            ws_risk.sheet_properties.tabColor = "FF9800"
            
            ws_risk.merge_cells('A1:H1')
            ws_risk['A1'] = '3. RISK ANALYSIS'
            ws_risk['A1'].font = subtitle_font
            
            if risk_data:
                risk_headers = ['Period', 'Year', 'Likelihood', 'Impact', 'Risk Score', 'Rating', 'Control Eff.', 'Residual Risk']
                for col_idx, text in enumerate(risk_headers, 1):
                    cell = ws_risk.cell(row=3, column=col_idx, value=text)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                
                for r_idx, r in enumerate(risk_data, 4):
                    vals = [
                        r.assessment_period, r.year, r.likelihood, r.impact,
                        r.risk_score, r.risk_rating.upper() if r.risk_rating else 'N/A',
                        r.control_effectiveness, f"{r.residual_risk:.2f}" if r.residual_risk else 'N/A'
                    ]
                    for c_idx, val in enumerate(vals, 1):
                        cell = ws_risk.cell(row=r_idx, column=c_idx, value=val)
                        cell.font = normal_font
                        cell.alignment = Alignment(horizontal='center')
                        cell.border = thin_border
            else:
                ws_risk['A3'] = 'No risk assessment data available for this engagement.'
                ws_risk['A3'].font = normal_font
                ws_risk.merge_cells('A3:H3')
            
            for col in range(1, 9):
                ws_risk.column_dimensions[get_column_letter(col)].width = 16
            
            # =========== SHEET 4: FINDINGS TABLE ===========
            ws_findings = wb.create_sheet("Findings")
            ws_findings.sheet_properties.tabColor = "F44336"
            
            ws_findings.merge_cells('A1:G1')
            ws_findings['A1'] = '4. FINDINGS SUMMARY TABLE'
            ws_findings['A1'].font = subtitle_font
            
            if findings:
                f_headers = ['#', 'Ref No.', 'Title', 'Severity', 'Category', 'Status', 'Recommendation']
                for col_idx, text in enumerate(f_headers, 1):
                    cell = ws_findings.cell(row=3, column=col_idx, value=text)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                
                for f_idx, f in enumerate(findings, 4):
                    vals = [
                        f_idx - 3,
                        f.finding_number,
                        f.title,
                        f.severity.upper() if f.severity else 'N/A',
                        f.category.replace('_', ' ').title() if f.category else 'N/A',
                        f.status.upper() if f.status else 'N/A',
                        f.recommendation or 'N/A',
                    ]
                    for c_idx, val in enumerate(vals, 1):
                        cell = ws_findings.cell(row=f_idx, column=c_idx, value=val)
                        cell.font = normal_font
                        cell.border = thin_border
                        if c_idx == 4:  # Severity column
                            cell.alignment = Alignment(horizontal='center')
                        if c_idx in (1, 5, 6):
                            cell.alignment = Alignment(horizontal='center')
                
                # Detailed findings section
                detail_row = len(findings) + 6
                ws_findings[f'A{detail_row}'] = 'Detailed Finding Descriptions:'
                ws_findings[f'A{detail_row}'].font = subtitle_font
                
                for f_idx, f in enumerate(findings):
                    dr = detail_row + 1 + (f_idx * 7)
                    ws_findings[f'A{dr}'] = f"{f.finding_number}: {f.title}"
                    ws_findings[f'A{dr}'].font = bold_font
                    ws_findings[f'C{dr+1}'] = f'Condition: {f.condition or "N/A"}'
                    ws_findings[f'A{dr+1}'].font = normal_font
                    ws_findings[f'C{dr+2}'] = f'Criteria: {f.criteria or "N/A"}'
                    ws_findings[f'A{dr+2}'].font = normal_font
                    ws_findings[f'C{dr+3}'] = f'Cause: {f.cause or "N/A"}'
                    ws_findings[f'A{dr+3}'].font = normal_font
                    ws_findings[f'C{dr+4}'] = f'Effect/Impact: {f.effect or "N/A"}'
                    ws_findings[f'A{dr+4}'].font = normal_font
                    ws_findings[f'C{dr+5}'] = f'Recommendation: {f.recommendation or "N/A"}'
                    ws_findings[f'A{dr+5}'].font = normal_font
                    ws_findings[f'C{dr+6}'] = f'Management Response: {f.management_response or "N/A"}'
                    ws_findings[f'A{dr+6}'].font = normal_font
            else:
                ws_findings['A3'] = 'No findings registered for this engagement.'
                ws_findings['A3'].font = normal_font
            
            ws_findings.column_dimensions['A'].width = 6
            ws_findings.column_dimensions['B'].width = 16
            ws_findings.column_dimensions['C'].width = 45
            ws_findings.column_dimensions['D'].width = 14
            ws_findings.column_dimensions['E'].width = 20
            ws_findings.column_dimensions['F'].width = 14
            ws_findings.column_dimensions['G'].width = 40
            
            # =========== SHEET 5: CAPA ===========
            ws_capa = wb.create_sheet("CAPA")
            ws_capa.sheet_properties.tabColor = "9C27B0"
            
            ws_capa.merge_cells('A1:H1')
            ws_capa['A1'] = '5. CORRECTIVE ACTION PLAN (CAPA)'
            ws_capa['A1'].font = subtitle_font
            
            if corrective_actions:
                capa_headers = ['#', 'Action No.', 'Title', 'Finding Ref', 'Priority', 'Status', 'Owner', 'Due Date']
                for col_idx, text in enumerate(capa_headers, 1):
                    cell = ws_capa.cell(row=3, column=col_idx, value=text)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = Alignment(horizontal='center')
                    cell.border = thin_border
                
                for ca_idx, ca in enumerate(corrective_actions, 4):
                    vals = [
                        ca_idx - 3,
                        ca.action_number,
                        ca.title,
                        ca.finding.finding_number if ca.finding else 'N/A',
                        ca.priority.upper() if ca.priority else 'N/A',
                        ca.status.upper() if ca.status else 'N/A',
                        ca.owner.full_name if ca.owner else 'Unassigned',
                        str(ca.due_date),
                    ]
                    for c_idx, val in enumerate(vals, 1):
                        cell = ws_capa.cell(row=ca_idx, column=c_idx, value=val)
                        cell.font = normal_font
                        cell.border = thin_border
                        if c_idx in (1, 4, 5, 6, 8):
                            cell.alignment = Alignment(horizontal='center')
                
                # Detailed CAPA section
                detail_row = len(corrective_actions) + 6
                ws_capa[f'A{detail_row}'] = 'Detailed Corrective Action Descriptions:'
                ws_capa[f'A{detail_row}'].font = subtitle_font
                
                for ca_idx, ca in enumerate(corrective_actions):
                    dr = detail_row + 1 + (ca_idx * 6)
                    ws_capa[f'A{dr}'] = f"{ca.action_number}: {ca.title}"
                    ws_capa[f'A{dr}'].font = bold_font
                    ws_capa[f'A{dr+1}'] = f"Finding: {ca.finding.finding_number + ' - ' + ca.finding.title if ca.finding else 'N/A'}"
                    ws_capa[f'A{dr+1}'].font = normal_font
                    ws_capa[f'A{dr+2}'] = f"Description: {ca.description or 'N/A'}"
                    ws_capa[f'A{dr+2}'].font = normal_font
                    ws_capa[f'A{dr+3}'] = f"Recommendation: {ca.recommendation or 'N/A'}"
                    ws_capa[f'A{dr+3}'].font = normal_font
                    ws_capa[f'A{dr+4}'] = f"Owner: {ca.owner.full_name if ca.owner else 'Unassigned'} | Due: {ca.due_date}"
                    ws_capa[f'A{dr+4}'].font = normal_font
                    if ca.management_response:
                        ws_capa[f'A{dr+5}'] = f"Management Response: {ca.management_response}"
                        ws_capa[f'A{dr+5}'].font = normal_font
            else:
                ws_capa['A3'] = 'No corrective actions have been assigned for findings in this engagement.'
                ws_capa['A3'].font = normal_font
                ws_capa.merge_cells('A3:H3')
            
            ws_capa.column_dimensions['A'].width = 6
            ws_capa.column_dimensions['B'].width = 18
            ws_capa.column_dimensions['C'].width = 45
            ws_capa.column_dimensions['D'].width = 16
            ws_capa.column_dimensions['E'].width = 12
            ws_capa.column_dimensions['F'].width = 16
            ws_capa.column_dimensions['G'].width = 25
            ws_capa.column_dimensions['H'].width = 16
            
            wb.save(buf)
            buf.seek(0)
            report.file.save(filename, ContentFile(buf.read()), save=False)
            report.status = 'ready'
            report.save()
            
            AuditTrail.objects.create(
                user=report.generated_by,
                action='EXPORT',
                model_name='GeneratedReport',
                object_id=str(report.id),
                object_repr=f"Generated Excel report: {report.title}",
                ip_address=None,
            )
            
        else:  # word or default text
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(10)
            
            # ========== COVER PAGE ==========
            for _ in range(4):
                doc.add_paragraph('')
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('ETHIOPIAN ELECTRIC UTILITY')
            run.bold = True
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('Internal Audit Department')
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            
            doc.add_paragraph('')
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(report.title)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            
            doc.add_paragraph('')
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Date: {timezone.now().strftime('%d %B %Y')}")
            run.font.size = Pt(11)
            
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Generated By: {report.generated_by.full_name if report.generated_by else 'System'}")
            run.font.size = Pt(11)
            
            doc.add_page_break()
            
            # ========== TABLE OF CONTENTS ==========
            p = doc.add_paragraph()
            run = p.add_run('Table of Contents')
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            
            toc_sections = ['1. Background', '2. Executive Summary', '3. Risk Analysis', '4. Findings Summary Table', '5. Corrective Action Plan (CAPA)']
            for item in toc_sections:
                p = doc.add_paragraph(item)
                p.paragraph_format.space_after = Pt(4)
            
            doc.add_page_break()
            
            # Helper function for section headers
            def add_section_header(text):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            
            def add_bold_text(text):
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
                return p
            
            def add_normal_text(text):
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(4)
                return p
            
            # ========== 1. BACKGROUND ==========
            add_section_header('1. Background')
            
            add_normal_text(
                f"This report presents the findings and recommendations from the audit engagement "
                f"{engagement_info.get('title', 'N/A')} conducted by the EEU Internal Audit Department. "
                f"The audit was performed in accordance with international standards."
            )
            
            # Background table
            table = doc.add_table(rows=8, cols=2)
            table.style = 'Light Grid Accent 1'
            bg_fields = [
                ('Engagement Title', engagement_info.get('title', 'N/A')),
                ('Engagement Type', engagement_info.get('type', 'N/A')),
                ('Department', engagement_info.get('department', 'N/A')),
                ('Lead Auditor', engagement_info.get('lead_auditor', 'N/A')),
                ('Supervisor', engagement_info.get('supervisor', 'N/A')),
                ('Risk Level', engagement_info.get('risk_level', 'N/A')),
                ('Planned Start', str(engagement_info.get('planned_start', 'N/A'))),
                ('Planned End', str(engagement_info.get('planned_end', 'N/A'))),
            ]
            for i, (label, value) in enumerate(bg_fields):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value
                for cell in table.rows[i].cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            doc.add_paragraph('')
            add_bold_text('Audit Objectives:')
            add_normal_text(engagement_info.get('objectives', 'No objectives defined.'))
            
            add_bold_text('Audit Scope:')
            add_normal_text(engagement_info.get('scope', 'No scope defined.'))
            
            doc.add_page_break()
            
            # ========== 2. EXECUTIVE SUMMARY ==========
            add_section_header('2. Executive Summary')
            
            add_normal_text(
                f"A total of {total_findings} finding(s) were identified during this audit engagement. "
                f"Of these, {critical_count} are Critical, {high_count} are High, "
                f"{medium_count} are Medium, and {low_count} are Low severity."
            )
            
            if total_findings > 0:
                table = doc.add_table(rows=len(severity_counts) + 1, cols=2)
                table.style = 'Light Grid Accent 1'
                table.rows[0].cells[0].text = 'Severity Level'
                table.rows[0].cells[1].text = 'Count'
                row_idx = 1
                for sev in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW', 'INFORMATIONAL']:
                    count = severity_counts.get(sev, 0)
                    if count > 0:
                        table.rows[row_idx].cells[0].text = sev
                        table.rows[row_idx].cells[1].text = str(count)
                        row_idx += 1
            
            doc.add_page_break()
            
            # ========== 3. RISK ANALYSIS ==========
            add_section_header('3. Risk Analysis')
            
            if risk_data:
                add_normal_text(f"Risk assessment data for {engagement_info.get('department', 'N/A')}:")
                table = doc.add_table(rows=min(len(risk_data), 15) + 1, cols=8)
                table.style = 'Light Grid Accent 1'
                headers = ['Period', 'Year', 'Likelihood', 'Impact', 'Risk Score', 'Rating', 'Ctrl Eff.', 'Residual']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for r_idx, r in enumerate(risk_data[:15]):
                    vals = [
                        r.assessment_period, str(r.year), str(r.likelihood), str(r.impact),
                        str(r.risk_score), r.risk_rating.upper() if r.risk_rating else 'N/A',
                        str(r.control_effectiveness), f"{r.residual_risk:.2f}" if r.residual_risk else 'N/A'
                    ]
                    for c_idx, val in enumerate(vals):
                        table.rows[r_idx + 1].cells[c_idx].text = val
            else:
                add_normal_text('No risk assessment data available for this engagement.')
            
            doc.add_page_break()
            
            # ========== 4. FINDINGS TABLE ==========
            add_section_header('4. Findings Summary Table')
            
            if findings:
                add_normal_text(f"The following table summarizes all {total_findings} finding(s):")
                table = doc.add_table(rows=len(findings) + 1, cols=7)
                table.style = 'Light Grid Accent 1'
                headers = ['#', 'Ref No.', 'Title', 'Severity', 'Category', 'Status', 'Recommendation']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for f_idx, f in enumerate(findings):
                    table.rows[f_idx + 1].cells[0].text = str(f_idx + 1)
                    table.rows[f_idx + 1].cells[1].text = f.finding_number
                    table.rows[f_idx + 1].cells[2].text = f.title[:60]
                    table.rows[f_idx + 1].cells[3].text = f.severity.upper() if f.severity else 'N/A'
                    table.rows[f_idx + 1].cells[4].text = f.category.replace('_', ' ').title() if f.category else 'N/A'
                    table.rows[f_idx + 1].cells[5].text = f.status.upper() if f.status else 'N/A'
                    table.rows[f_idx + 1].cells[6].text = f.recommendation[:80] if f.recommendation else 'N/A'
                
                # Detailed findings
                doc.add_paragraph('')
                add_bold_text('Detailed Finding Descriptions:')
                for f in findings:
                    add_bold_text(f"{f.finding_number}: {f.title}")
                    add_normal_text(f"Condition: {f.condition or 'N/A'}")
                    add_normal_text(f"Criteria: {f.criteria or 'N/A'}")
                    add_normal_text(f"Cause: {f.cause or 'N/A'}")
                    add_normal_text(f"Effect/Impact: {f.effect or 'N/A'}")
                    add_normal_text(f"Recommendation: {f.recommendation or 'N/A'}")
                    add_normal_text(f"Management Response: {f.management_response or 'N/A'}")
                    doc.add_paragraph('')
            else:
                add_normal_text('No findings were registered for this engagement.')
            
            doc.add_page_break()
            
            # ========== 5. CAPA ==========
            add_section_header('5. Corrective Action Plan (CAPA)')
            
            if corrective_actions:
                add_normal_text(f"The following {len(corrective_actions)} corrective action(s) have been defined:")
                table = doc.add_table(rows=len(corrective_actions) + 1, cols=8)
                table.style = 'Light Grid Accent 1'
                headers = ['#', 'Action No.', 'Title', 'Finding Ref', 'Priority', 'Status', 'Owner', 'Due Date']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                for ca_idx, ca in enumerate(corrective_actions):
                    table.rows[ca_idx + 1].cells[0].text = str(ca_idx + 1)
                    table.rows[ca_idx + 1].cells[1].text = ca.action_number
                    table.rows[ca_idx + 1].cells[2].text = ca.title[:50]
                    table.rows[ca_idx + 1].cells[3].text = ca.finding.finding_number if ca.finding else 'N/A'
                    table.rows[ca_idx + 1].cells[4].text = ca.priority.upper() if ca.priority else 'N/A'
                    table.rows[ca_idx + 1].cells[5].text = ca.status.upper() if ca.status else 'N/A'
                    table.rows[ca_idx + 1].cells[6].text = ca.owner.full_name if ca.owner else 'Unassigned'
                    table.rows[ca_idx + 1].cells[7].text = str(ca.due_date)
                
                # Detailed CAPA
                doc.add_paragraph('')
                add_bold_text('Detailed Corrective Action Descriptions:')
                for ca in corrective_actions:
                    add_bold_text(f"{ca.action_number}: {ca.title}")
                    add_normal_text(f"Finding: {ca.finding.finding_number + ' - ' + ca.finding.title if ca.finding else 'N/A'}")
                    add_normal_text(f"Description: {ca.description or 'N/A'}")
                    add_normal_text(f"Recommendation: {ca.recommendation or 'N/A'}")
                    add_normal_text(f"Owner: {ca.owner.full_name if ca.owner else 'Unassigned'} | Due: {ca.due_date}")
                    if ca.management_response:
                        add_normal_text(f"Management Response: {ca.management_response}")
                    doc.add_paragraph('')
            else:
                add_normal_text('No corrective actions have been assigned for findings in this engagement.')
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            report.file.save(filename, ContentFile(buf.read()), save=False)
            report.status = 'ready'
            report.save()
            
            AuditTrail.objects.create(
                user=report.generated_by,
                action='EXPORT',
                model_name='GeneratedReport',
                object_id=str(report.id),
                object_repr=f"Generated Word report: {report.title}",
                ip_address=None,
            )

    @action(detail=True, methods=['get'], url_path='export')
    def export(self, request, pk=None):
        report = self.get_object()
        if report.file:
            response = HttpResponse(report.file.read(), content_type='application/octet-stream')
            response['Content-Disposition'] = f'attachment; filename="{report.file.name.split("/")[-1]}"'
            return response
        return Response({'detail': 'Report file not generated yet.'}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=False, methods=['post'], url_path='generate-pdf')
    def generate_pdf(self, request):
        """Generate a PDF report using ReportLab"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        elements = []

        title = request.data.get('title', 'EEU Internal Audit Report')
        elements.append(Paragraph(title, styles['Title']))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(
            f"Generated: {timezone.now().strftime('%d %B %Y %H:%M')}",
            styles['Normal']
        ))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(
            "Ethiopian Electric Utility — Internal Audit Department",
            styles['Normal']
        ))

        content = request.data.get('content', '')
        if content:
            elements.append(Spacer(1, 12))
            elements.append(Paragraph(content, styles['Normal']))

        doc.build(elements)
        buf.seek(0)

        report = GeneratedReport.objects.create(
            title=title, format='pdf', status='ready',
            generated_by=request.user,
            parameters=request.data
        )

        response = HttpResponse(buf.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{title}.pdf"'
        return response

    @action(detail=False, methods=['post'], url_path='generate-excel')
    def generate_excel(self, request):
        """Generate an Excel report using openpyxl"""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment

        wb = openpyxl.Workbook()
        ws = wb.active
        title = request.data.get('title', 'EEU Audit Report')
        ws.title = 'Audit Report'

        # Header
        ws['A1'] = 'Ethiopian Electric Utility — Internal Audit Department'
        ws['A1'].font = Font(bold=True, size=14)
        ws['A2'] = title
        ws['A2'].font = Font(bold=True, size=12)
        ws['A3'] = f"Generated: {timezone.now().strftime('%d/%m/%Y %H:%M')}"
        ws.column_dimensions['A'].width = 40

        headers = request.data.get('headers', ['Item', 'Details'])
        rows = request.data.get('rows', [])
        start_row = 5
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=start_row, column=col, value=h)
            cell.font = Font(bold=True, color='FFFFFF')
            cell.fill = PatternFill(fill_type='solid', fgColor='1E3A5F')
        for r_idx, row in enumerate(rows, start_row + 1):
            for c_idx, val in enumerate(row, 1):
                ws.cell(row=r_idx, column=c_idx, value=val)

        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)

        response = HttpResponse(
            buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{title}.xlsx"'
        return response

    @action(detail=False, methods=['get'], url_path='analytics')
    def analytics(self, request):
        """Return aggregated analytics data for dashboards"""
        from apps.findings.models import AuditFinding
        from apps.corrective_actions.models import CorrectiveAction
        from apps.audit_planning.models import AuditEngagement
        from django.db.models import Count
        from django.utils import timezone
        import datetime

        today = timezone.now().date()
        six_months = [(today.replace(day=1) - datetime.timedelta(days=30 * i)) for i in range(5, -1, -1)]

        monthly_findings = []
        for m in six_months:
            count = AuditFinding.objects.filter(
                created_at__year=m.year, created_at__month=m.month
            ).count()
            monthly_findings.append({'month': m.strftime('%b %Y'), 'count': count})

        return Response({
            'findings_by_severity': list(
                AuditFinding.objects.values('severity').annotate(count=Count('id'))
            ),
            'findings_by_category': list(
                AuditFinding.objects.values('category').annotate(count=Count('id'))
            ),
            'actions_by_status': list(
                CorrectiveAction.objects.values('status').annotate(count=Count('id'))
            ),
            'engagements_by_type': list(
                AuditEngagement.objects.values('engagement_type').annotate(count=Count('id'))
            ),
            'monthly_findings': monthly_findings,
        })